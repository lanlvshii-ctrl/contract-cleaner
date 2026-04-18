#!/usr/bin/env python3
"""
docx_exporter.py - 清洗完成后生成 Word 交付物

三个输出：
  1. 合同-原始版2.0.docx   — 原始MD用 md2docx 引擎转成美观 Word
  2. 合同-清洁版.docx      — 清洗后MD用 md2docx 引擎转成美观 Word
  3. 合同-修订痕迹版.docx  — 预处理两份MD后用 pandiff 对比生成修订标记

修订痕迹版流程：
  原始MD → 预处理（清反斜杠+清列表语义+空行标准化）┐
                                                    ├→ pandiff → 清除自动编号 → 修订痕迹版.docx
  清洗MD → 预处理（空行标准化）────────────────────┘

依赖：
  - python-docx (pip install python-docx)
  - pandiff (npm install -g pandiff)
  - pandoc (brew install pandoc) — pandiff 依赖
"""

import logging
import subprocess
import shutil
import os
import re
import tempfile
from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

logger = logging.getLogger('contract_cleaner')


# ============================================
# 排版引擎（源自 md2docx_plain.py）
# ============================================

def clean_markdown_formatting(text):
    """清除Markdown格式符号，保留加粗标记用于后续处理，清除斜体"""
    bold_placeholders = []
    def save_bold(match):
        bold_placeholders.append(match.group(1))
        return f"\x00BOLD{len(bold_placeholders)-1}\x00"
    
    text = re.sub(r'\*\*([^*]+)\*\*', save_bold, text)
    
    # 清除斜体标记
    text = re.sub(r'\*([^*]+)\*', r'\1', text)
    text = re.sub(r'_([^_]+)_', r'\1', text)
    
    # 恢复加粗占位符
    for i, content in enumerate(bold_placeholders):
        text = text.replace(f"\x00BOLD{i}\x00", f"**{content}**")
    
    # 清除行内代码、链接、图片
    text = re.sub(r'`([^`]+)`', r'\1', text)
    text = re.sub(r'\[([^\]]+)\]\([^\)]+\)', r'\1', text)
    text = re.sub(r'!\[([^\]]*)\]\([^\)]+\)', r'【\1】', text)
    text = text.replace('[', '【').replace(']', '】')
    
    # 清除反斜杠残留（用户反馈：md2docx 后文档中有很多 \ 符号）
    text = text.replace('\\', '')
    
    return text


def add_formatted_text(paragraph, text, font_name='微软雅黑', font_size=Pt(10),
                       bold=False, color=RGBColor(0x00, 0x00, 0x00)):
    """向段落添加文本，处理 **加粗** 标记"""
    parts = re.split(r'(\*\*[^*]+\*\*)', text)
    for part in parts:
        if part.startswith('**') and part.endswith('**'):
            run = paragraph.add_run(part[2:-2])
            run.font.name = font_name
            run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
            run.font.size = font_size
            run.font.bold = True
            run.font.color.rgb = color
        else:
            if part:
                run = paragraph.add_run(part)
                run.font.name = font_name
                run._element.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                run.font.size = font_size
                run.font.bold = bold
                run.font.color.rgb = color


def set_cell_shading(cell, fill_color):
    """设置单元格背景色"""
    shading_elm = OxmlElement('w:shd')
    shading_elm.set(qn('w:fill'), fill_color)
    cell._tc.get_or_add_tcPr().append(shading_elm)


def is_separator_row(cells):
    """判断是否为Markdown表格分隔行"""
    if not cells:
        return False
    separator_pattern = re.compile(r'^[\s\-:]+$')
    return all(separator_pattern.match(cell) for cell in cells)


def parse_table_line(line):
    """解析表格行"""
    temp_placeholder = '\x00PIPE\x00'
    line = line.replace('\\|', temp_placeholder)
    cells = [cell.strip() for cell in line.split('|')[1:-1]]
    cells = [cell.replace(temp_placeholder, '|') for cell in cells]
    return cells


def create_beautiful_table(doc, table_data):
    """创建美观的表格"""
    if len(table_data) < 1:
        return
    
    max_cols = max(len(row) for row in table_data)
    if max_cols == 0:
        return
    
    # 统一列数
    normalized_data = []
    for row in table_data:
        if len(row) < max_cols:
            row = row + [''] * (max_cols - len(row))
        elif len(row) > max_cols:
            row = row[:max_cols]
        normalized_data.append(row)
    
    # 清除Markdown标记
    cleaned_data = []
    for row in normalized_data:
        cleaned_row = [clean_markdown_formatting(cell) for cell in row]
        cleaned_data.append(cleaned_row)
    
    table = doc.add_table(rows=len(cleaned_data), cols=max_cols)
    table.style = 'Table Grid'
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.allow_autofit = False
    
    for i, row_data in enumerate(cleaned_data):
        row = table.rows[i]
        for j, cell_text in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ''
            
            if '**' in cell_text:
                parts = re.split(r'(\*\*[^*]+\*\*)', cell_text)
                for part in parts:
                    if part.startswith('**') and part.endswith('**'):
                        run = cell.paragraphs[0].add_run(part[2:-2])
                        run.font.name = '微软雅黑'
                        run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                        run.font.size = Pt(10)
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    else:
                        if part:
                            run = cell.paragraphs[0].add_run(part)
                            run.font.name = '微软雅黑'
                            run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                            run.font.size = Pt(10)
                            run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
            else:
                cell.text = cell_text
            
            for paragraph in cell.paragraphs:
                paragraph.alignment = WD_ALIGN_PARAGRAPH.LEFT
                for run in paragraph.runs:
                    run.font.name = '微软雅黑'
                    run._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
                    run.font.size = Pt(10)
                    if i == 0:
                        run.font.bold = True
                        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                        paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
    
    # 表头背景色
    header_row = table.rows[0]
    for cell in header_row.cells:
        set_cell_shading(cell, 'D3D3D3')
    
    for row in table.rows:
        row.height = Pt(30)
    
    doc.add_paragraph()


def md_to_docx(md_path: Path, output_path: Path) -> bool:
    """
    使用 md2docx 排版引擎将 MD 转为美观的 docx。
    复刻自 md2docx_plain.py 的排版逻辑。
    """
    if not md_path.exists():
        logger.error(f'MD文件不存在: {md_path}')
        return False
    
    content = md_path.read_text(encoding='utf-8')
    lines = content.split('\n')
    
    doc = Document()
    
    # 设置文档默认字体
    style = doc.styles['Normal']
    style.font.name = '微软雅黑'
    style._element.rPr.rFonts.set(qn('w:eastAsia'), '微软雅黑')
    style.font.size = Pt(10)
    style.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
    
    # 清除默认样式的段落编号
    numPr = style._element.find(qn('w:numPr'))
    if numPr is not None:
        style._element.remove(numPr)
    
    in_code_block = False
    code_block_content = []
    in_table = False
    table_data = []
    
    i = 0
    while i < len(lines):
        line = lines[i].rstrip()
        stripped_line = line.lstrip()
        
        # 代码块
        if stripped_line.startswith('```'):
            if in_code_block:
                if code_block_content:
                    code_para = doc.add_paragraph()
                    code_para.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
                    code_para.paragraph_format.left_indent = Inches(0)
                    code_para.paragraph_format.first_line_indent = Inches(0)
                    for code_line in code_block_content:
                        clean_code = clean_markdown_formatting(code_line)
                        run = code_para.add_run(clean_code + '\n')
                        run.font.name = 'Consolas'
                        run.font.size = Pt(10)
                        run.font.color.rgb = RGBColor(0x00, 0x00, 0x00)
                    code_block_content = []
                in_code_block = False
            else:
                in_code_block = True
            i += 1
            continue
        
        if in_code_block:
            code_block_content.append(line)
            i += 1
            continue
        
        # 表格
        if stripped_line.startswith('|') and not in_table:
            in_table = True
            table_data = []
        
        if in_table:
            if stripped_line.startswith('|'):
                cells = parse_table_line(stripped_line)
                if cells and not is_separator_row(cells):
                    table_data.append(cells)
                i += 1
            elif stripped_line == '':
                if i + 1 < len(lines) and lines[i + 1].lstrip().startswith('|'):
                    i += 1
                    continue
                else:
                    if len(table_data) >= 1:
                        create_beautiful_table(doc, table_data)
                    in_table = False
                    table_data = []
            else:
                if len(table_data) >= 1:
                    create_beautiful_table(doc, table_data)
                in_table = False
                table_data = []
        else:
            if not line:
                i += 1
                continue
            
            # 分割线
            if stripped_line in ('---', '***', '___'):
                i += 1
                continue
            
            # 标题
            if line.startswith('# ') and not line.startswith('## '):
                text = clean_markdown_formatting(line[2:])
                if text:
                    p = doc.add_paragraph()
                    add_formatted_text(p, text, bold=True)
                    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    doc.add_paragraph()
            elif line.startswith('## '):
                text = clean_markdown_formatting(line[3:])
                if text:
                    p = doc.add_paragraph()
                    add_formatted_text(p, text, bold=True)
            elif line.startswith('### '):
                text = clean_markdown_formatting(line[4:])
                if text:
                    p = doc.add_paragraph()
                    add_formatted_text(p, text, bold=True)
            elif line.startswith('#### '):
                text = clean_markdown_formatting(line[5:])
                if text:
                    p = doc.add_paragraph()
                    add_formatted_text(p, text, bold=True)
            elif line.startswith('- ') or line.startswith('* '):
                text = clean_markdown_formatting(line)
                if text:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0)
                    p.paragraph_format.first_line_indent = Inches(0)
                    add_formatted_text(p, text)
            elif re.match(r'^\d+\.\s+', line):
                text = clean_markdown_formatting(line)
                if text:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0)
                    p.paragraph_format.first_line_indent = Inches(0)
                    add_formatted_text(p, text)
            elif line.startswith('> '):
                text = clean_markdown_formatting(line[2:])
                if text:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0)
                    p.paragraph_format.first_line_indent = Inches(0)
                    add_formatted_text(p, text)
            else:
                text = clean_markdown_formatting(line)
                if text:
                    p = doc.add_paragraph()
                    p.paragraph_format.left_indent = Inches(0)
                    p.paragraph_format.first_line_indent = Inches(0)
                    add_formatted_text(p, text)
        
        i += 1
    
    # 最后一个表格
    if in_table and len(table_data) >= 1:
        create_beautiful_table(doc, table_data)
    
    # 清除所有段落的自动编号
    _remove_all_numbering(doc)
    
    # 保存
    _save_docx_safely(doc, output_path)
    return True


def _remove_all_numbering(doc: Document):
    """清除文档中所有段落的自动编号，并重置列表样式为 Normal"""
    normal_style = doc.styles['Normal']
    for paragraph in doc.paragraphs:
        numPr = paragraph._element.find(qn('w:numPr'))
        if numPr is not None:
            paragraph._element.remove(numPr)
        # 重置列表样式（pandiff 可能通过样式应用编号）
        if paragraph.style and paragraph.style.name and 'List' in paragraph.style.name:
            paragraph.style = normal_style
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    numPr = paragraph._element.find(qn('w:numPr'))
                    if numPr is not None:
                        paragraph._element.remove(numPr)
                    if paragraph.style and paragraph.style.name and 'List' in paragraph.style.name:
                        paragraph.style = normal_style


def _save_docx_safely(doc: Document, output_path: Path):
    """安全保存docx，移除macOS隔离属性"""
    with tempfile.NamedTemporaryFile(suffix='.docx', delete=False) as tmp_file:
        tmp_path = tmp_file.name
    
    doc.save(tmp_path)
    
    try:
        result = subprocess.run(['cp', '-X', tmp_path, str(output_path)],
                               capture_output=True, text=True, timeout=5)
        if result.returncode != 0:
            with open(tmp_path, 'rb') as src, open(str(output_path), 'wb') as dst:
                dst.write(src.read())
        
        for attr in ['com.apple.quarantine', 'com.apple.provenance',
                      'com.apple.macl', 'com.apple.metadata:kMDItemWhereFroms']:
            try:
                subprocess.run(['xattr', '-d', attr, str(output_path)],
                              capture_output=True, check=False, timeout=2)
            except Exception:
                pass
        
        os.chmod(str(output_path), 0o644)
    except Exception as e:
        logger.warning(f'文件保存异常: {e}')
        os.replace(tmp_path, str(output_path))
    finally:
        try:
            if os.path.exists(tmp_path):
                os.unlink(tmp_path)
        except Exception:
            pass


# ============================================
# Word 内置对比功能
# ============================================

def _preprocess_md(text: str, *, is_old: bool = False) -> str:
    """
    预处理 Markdown 文本，为 pandiff 对比做准备。
    复刻自 md_compare_docx.py 的预处理逻辑。

    处理项：
    1. 清除反斜杠残留（docx→md 转换产生的 \\ 符号）— 仅旧文件
    2. 清除有序列表语义（防止 Word 重新生成自动编号）
    3. 删除无序列表标记 "- "，转为普通段落 — 仅旧文件
    4. 段落间标准化空行（每段之间确保恰好一个空行）
    """
    # ---- 1. 清除反斜杠（仅旧文件）----
    if is_old:
        text = text.replace('\\', '')

    # ---- 2. 清除有序列表语义 ----
    # 将 "3. 内容" / "> 3. 内容" 转义为 "3\. 内容"，防止 pandoc 识别为有序列表生成自动编号
    lines = text.split('\n')
    result = []
    for line in lines:
        m = re.match(r'^((?:>\s*)*)( {0,3})(\d+)([.)]) (.*)$', line)
        if m:
            quote_prefix, indent, num, sep, content = m.groups()
            result.append(f"{quote_prefix}{indent}{num}\\{sep} {content}")
        else:
            result.append(line)

    text = '\n'.join(result)

    # 清除 pandoc 编号重置注释 <!-- -->
    text = re.sub(r'\n?<!--\s*-->\n?', '\n', text)

    # ---- 3. 删除无序列表标记（仅旧文件）----
    if is_old:
        lines = text.split('\n')
        cleaned = []
        for line in lines:
            m = re.match(r'^((?:>\s*)*)( {0,3})- (.+)$', line)
            if m:
                quote_prefix, indent, content = m.groups()
                cleaned.append(f"{quote_prefix}{indent}{content}")
            else:
                cleaned.append(line)
        text = '\n'.join(cleaned)

    # ---- 4. 段落间标准化空行 ----
    lines = text.split('\n')
    normalized = []
    i = 0
    while i < len(lines):
        line = lines[i]
        normalized.append(line)
        if line.strip() and i < len(lines) - 1:
            blank_count = 0
            j = i + 1
            while j < len(lines) and lines[j].strip() == '':
                blank_count += 1
                j += 1
            if j < len(lines) and lines[j].strip():
                normalized.append('')
                i = j
            else:
                i = j
        else:
            i += 1

    text = '\n'.join(normalized)
    text = text.rstrip('\n') + '\n'

    return text


def _pandiff_compare_md(old_md_path: Path, new_md_path: Path, output_docx: Path) -> bool:
    """
    用 pandiff 对比两个 MD 文件，生成带修订痕迹的 docx。

    输入必须是已预处理的交付物 MD 文件，不再内部做预处理。
    流程：pandiff 对比 → 清除自动编号 → 保存
    """
    if not shutil.which('pandiff'):
        logger.warning('  pandiff 未安装，跳过修订痕迹版')
        return False

    with tempfile.TemporaryDirectory() as tmp_dir:
        tmp = Path(tmp_dir)
        temp_output = tmp / 'redline.docx'
        try:
            result = subprocess.run(
                ['pandiff', str(old_md_path), str(new_md_path), '-o', str(temp_output)],
                capture_output=True, text=True, timeout=120
            )
            if result.returncode != 0 or not temp_output.exists():
                logger.warning(f'  pandiff 对比失败: {result.stderr.strip()[:200]}')
                return False
        except (FileNotFoundError, subprocess.TimeoutExpired) as e:
            logger.warning(f'  pandiff 执行异常: {e}')
            return False

        # 后处理：清除自动编号
        try:
            doc = Document(str(temp_output))
            _remove_all_numbering(doc)
            _save_docx_safely(doc, output_docx)
            return True
        except Exception as e:
            # 后处理失败则直接复制
            shutil.copy2(str(temp_output), str(output_docx))
            return True


# ============================================
# 导出接口
# ============================================

def export_preprocessed_old_md(original_md: Path, output_dir: Path, stem: str) -> Path | None:
    """
    输出原合同（预处理后）.md。

    内容直接取自 _原始轻量.md，不再做额外预处理，确保用户用其他工具对比时基线一致。
    """
    output_path = output_dir / f'{stem}-原合同（预处理后）.md'

    try:
        text = original_md.read_text(encoding='utf-8')
        output_path.write_text(text, encoding='utf-8')
        logger.info(f'  📄 {output_path}')
        return output_path
    except Exception as e:
        logger.warning(f'  ⚠️ 原合同预处理 MD 输出出错: {e}')
        return None


def export_preprocessed_new_md(cleaned_md: Path, output_dir: Path, stem: str) -> Path | None:
    """
    一项预处理清洗版 MD 并输出为交付物。
    预处理：空行标准化（保证与原合同 MD 格式对称）。
    """
    output_path = output_dir / f'{stem}-新合同（预处理后）.md'

    try:
        text = cleaned_md.read_text(encoding='utf-8')
        preprocessed = _preprocess_md(text, is_old=False)
        output_path.write_text(preprocessed, encoding='utf-8')
        logger.info(f'  📄 {output_path}')
        return output_path
    except Exception as e:
        logger.warning(f'  ⚠️ 新合同预处理 MD 输出出错: {e}')
        return None


def export_clean_docx(cleaned_md: Path, output_dir: Path, stem: str) -> Path | None:
    """
    用 md2docx 排版引擎将清洗后 MD 转为美观的 docx。
    """
    output_docx = output_dir / f'{stem}-清洁版.docx'

    try:
        if md_to_docx(cleaned_md, output_docx):
            logger.info(f'  📄 {output_docx}')
            return output_docx
        else:
            logger.warning('  ⚠️ 清洁版 docx 生成失败')
            return None
    except Exception as e:
        logger.warning(f'  ⚠️ 清洁版 docx 生成出错: {e}')
        return None


def export_redline_docx(preprocessed_old_md: Path, preprocessed_new_md: Path, output_dir: Path, stem: str) -> Path | None:
    """
    用 pandiff 对比两个预处理后的 MD 文件，生成带修订痕迹的 docx。

    输入就是交付物的两个预处理 MD，不再内部做预处理。

    Args:
        preprocessed_old_md: 预处理后的原合同 MD（对比基准 = 前）
        preprocessed_new_md: 预处理后的新合同 MD（修订目标 = 后）
        output_dir:         输出目录
        stem:               文件名前缀
    """
    output_docx = output_dir / f'{stem}-对比版.docx'

    if not preprocessed_old_md or not preprocessed_old_md.exists():
        logger.warning('  ⚠️ 预处理后的原合同 MD 不存在，跳过对比版')
        return None

    if not preprocessed_new_md or not preprocessed_new_md.exists():
        logger.warning('  ⚠️ 预处理后的新合同 MD 不存在，跳过对比版')
        return None

    try:
        if _pandiff_compare_md(preprocessed_old_md, preprocessed_new_md, output_docx):
            logger.info(f'  📝 {output_docx}')
            return output_docx
        else:
            logger.warning('  ⚠️ 对比版 docx 生成失败（pandiff 不可用）')
            return None
    except Exception as e:
        logger.warning(f'  ⚠️ 对比版 docx 生成出错: {e}')
        return None


def export_docx_outputs(
    original_md: Path,
    cleaned_md: Path,
    output_dir: Path,
    stem: str,
) -> dict:
    """
    生成全部交付物。

    流程：
      原始MD → 三项预处理 → 原合同（预处理后）.md
      清洗MD → 一项预处理 → 新合同（预处理后）.md
      清洗MD → md2docx → 清洁版.docx
      预处理后原合同MD + 预处理后新合同MD → pandiff → 对比版.docx

    Args:
        original_md: 格式清洗后（AI 清洗前）的 MD
        cleaned_md:  最终清洗结果 MD
        output_dir: 输出目录
        stem:        原始文件名（不含扩展名）

    Returns:
        {'preprocessed_old': Path | None, 'preprocessed_new': Path | None,
         'clean': Path | None, 'redline': Path | None}
    """
    logger.info('\n【步骤5】生成交付物')

    # 1. 原合同（预处理后）MD
    pp_old_md = export_preprocessed_old_md(original_md, output_dir, stem)

    # 2. 新合同（预处理后）MD
    pp_new_md = export_preprocessed_new_md(cleaned_md, output_dir, stem)

    # 3. 清洁版 docx
    clean_docx = export_clean_docx(cleaned_md, output_dir, stem)

    # 4. 对比版 docx（pandiff 用交付物的两个 MD）
    redline_docx = None
    if pp_old_md and pp_new_md:
        redline_docx = export_redline_docx(pp_old_md, pp_new_md, output_dir, stem)

    if not pp_old_md and not pp_new_md and not clean_docx and not redline_docx:
        logger.info('  （未生成任何交付物，请检查依赖是否已安装）')

    return {
        'preprocessed_old': pp_old_md,
        'preprocessed_new': pp_new_md,
        'clean': clean_docx,
        'redline': redline_docx,
    }
